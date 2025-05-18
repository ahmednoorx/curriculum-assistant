from docx import Document

class DOCXExporter:
    def export(self, content, filename):
        doc = Document()
        doc.add_heading('Curriculum Assistant Export', level=1)
        doc.add_paragraph(content)
        doc.save(filename)